"""
DocuMind AI - Document Processor

Handles loading and chunking of PDF and DOCX documents.
"""

import re
from typing import List, Optional
from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document

# Chunking configuration
CHUNK_SIZE = 500  # tokens
CHUNK_OVERLAP = 50  # tokens


class DocumentProcessor:
    """Processes documents for embedding and storage."""

    @staticmethod
    def load_document(filepath: str) -> str:
        """Load a document and return its text content."""
        path = Path(filepath)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return DocumentProcessor._load_pdf(filepath)
        elif suffix == ".docx":
            return DocumentProcessor._load_docx(filepath)
        elif suffix == ".txt":
            return DocumentProcessor._load_txt(filepath)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    @staticmethod
    def _load_pdf(filepath: str) -> str:
        """Extract text from a PDF file."""
        try:
            reader = PdfReader(filepath)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)
        except Exception as e:
            raise RuntimeError(f"Error loading PDF: {e}")

    @staticmethod
    def _load_docx(filepath: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = Document(filepath)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n\n".join(text_parts)
        except Exception as e:
            raise RuntimeError(f"Error loading DOCX: {e}")

    @staticmethod
    def _load_txt(filepath: str) -> str:
        """Load a text file."""
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            raise RuntimeError(f"Error loading text file: {e}")

    @staticmethod
    def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
        """Split text into overlapping chunks for embedding."""
        # Simple character-based chunking as a approximation of token chunking
        # For production, use tiktoken or similar for accurate token counting

        # Clean the text
        text = DocumentProcessor._clean_text(text)

        if not text:
            return []

        # Split by sentences first to preserve semantic units
        sentences = re.split(r'(?<=[.!?])\s+', text)

        chunks = []
        current_chunk = []
        current_size = 0

        for sentence in sentences:
            sentence_size = len(sentence.split())

            # If single sentence exceeds chunk size, split it by words
            if sentence_size > chunk_size:
                if current_chunk:
                    chunks.append(" ".join(current_chunk))
                    current_chunk = current_chunk[-overlap:] if len(current_chunk) > overlap else []
                    current_size = sum(len(s.split()) for s in current_chunk)

                # Split long sentence
                words = sentence.split()
                for i in range(0, len(words), chunk_size - overlap):
                    chunk_words = words[i:i + chunk_size]
                    if i > 0:
                        # Add overlap from previous chunk
                        prev_words = words[max(0, i - overlap):i]
                        chunks.append(" ".join(prev_words + chunk_words[:overlap]))
                    else:
                        chunks.append(" ".join(chunk_words))

                current_chunk = []
                current_size = 0
                continue

            # Check if adding this sentence would exceed chunk size
            if current_size + sentence_size > chunk_size and current_chunk:
                # Save current chunk
                chunks.append(" ".join(current_chunk))

                # Start new chunk with overlap
                overlap_count = min(overlap, len(current_chunk))
                current_chunk = current_chunk[-overlap_count:] if overlap_count > 0 else []
                current_size = sum(len(s.split()) for s in current_chunk)

            current_chunk.append(sentence)
            current_size += sentence_size

        # Don't forget the last chunk
        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,!?;:\'"()\-–—]', '', text)
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        return text.strip()

    @staticmethod
    def process_document(filepath: str) -> List[str]:
        """Load a document and return its chunks."""
        text = DocumentProcessor.load_document(filepath)
        return DocumentProcessor.chunk_text(text)

    @staticmethod
    def get_document_summary(text: str, num_lines: int = 3) -> str:
        """Get a brief summary of the document (first few sentences)."""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return " ".join(sentences[:num_lines])
